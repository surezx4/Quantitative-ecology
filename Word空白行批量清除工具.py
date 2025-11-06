import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import threading

class WordBlankLineRemover:
    def __init__(self, root):
        self.root = root
        self.root.title("Word空白行批量清除工具")
        self.root.geometry("700x500")
        
        # 变量初始化
        self.input_folder = tk.StringVar()
        self.output_folder = tk.StringVar()
        
        # 选项变量
        self.remove_paragraph_blanks = tk.BooleanVar(value=True)
        self.remove_table_blanks = tk.BooleanVar(value=True)
        self.remove_section_blanks = tk.BooleanVar(value=True)
        
        self.setup_ui()
    
    def setup_ui(self):
        # 输入文件夹选择
        input_frame = tk.Frame(self.root)
        input_frame.pack(pady=10, padx=10, fill=tk.X)
        
        tk.Label(input_frame, text="输入文件夹:").grid(row=0, column=0, sticky=tk.W)
        tk.Entry(input_frame, textvariable=self.input_folder, width=50).grid(row=0, column=1, padx=5)
        tk.Button(input_frame, text="浏览", command=self.select_input_folder).grid(row=0, column=2)
        
        # 输出文件夹选择
        output_frame = tk.Frame(self.root)
        output_frame.pack(pady=10, padx=10, fill=tk.X)
        
        tk.Label(output_frame, text="输出文件夹:").grid(row=0, column=0, sticky=tk.W)
        tk.Entry(output_frame, textvariable=self.output_folder, width=50).grid(row=0, column=1, padx=5)
        tk.Button(output_frame, text="浏览", command=self.select_output_folder).grid(row=0, column=2)
        
        # 选项框架
        options_frame = tk.LabelFrame(self.root, text="处理选项")
        options_frame.pack(pady=10, padx=10, fill=tk.X)
        
        tk.Checkbutton(options_frame, text="删除段落中的空白行", 
                      variable=self.remove_paragraph_blanks).grid(row=0, column=0, sticky=tk.W, padx=5)
        tk.Checkbutton(options_frame, text="删除表格中的空白行", 
                      variable=self.remove_table_blanks).grid(row=0, column=1, sticky=tk.W, padx=5)
        tk.Checkbutton(options_frame, text="删除节中的空白段落", 
                      variable=self.remove_section_blanks).grid(row=0, column=2, sticky=tk.W, padx=5)
        
        # 处理按钮
        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=20)
        
        tk.Button(button_frame, text="开始处理", command=self.start_processing, 
                 bg="lightblue", font=("Arial", 12)).pack(pady=5)
        
        # 进度条
        progress_frame = tk.Frame(self.root)
        progress_frame.pack(pady=10, padx=10, fill=tk.X)
        
        self.progress = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, 
                                       length=500, mode='determinate')
        self.progress.pack(pady=5)
        
        self.status_label = tk.Label(progress_frame, text="准备就绪")
        self.status_label.pack()
        
        # 结果列表框
        result_frame = tk.Frame(self.root)
        result_frame.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)
        
        tk.Label(result_frame, text="处理结果:").pack(anchor=tk.W)
        
        self.result_text = tk.Text(result_frame, height=15)
        self.result_text.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 添加滚动条
        scrollbar = tk.Scrollbar(self.result_text)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.result_text.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.result_text.yview)
    
    def select_input_folder(self):
        folder = filedialog.askdirectory(title="选择包含Word文件的文件夹")
        if folder:
            self.input_folder.set(folder)
    
    def select_output_folder(self):
        folder = filedialog.askdirectory(title="选择输出文件夹")
        if folder:
            self.output_folder.set(folder)
    
    def is_blank_paragraph(self, paragraph):
        """检查段落是否为空行"""
        if not paragraph.text.strip():
            return True
        return False
    
    def remove_paragraph_blank_lines(self, doc):
        """从Word文档中移除段落空白行"""
        # 从后向前遍历，避免删除时索引变化问题
        for i in range(len(doc.paragraphs) - 1, -1, -1):
            if self.is_blank_paragraph(doc.paragraphs[i]):
                p = doc.paragraphs[i]._element
                p.getparent().remove(p)
    
    def remove_table_blank_lines(self, doc):
        """从Word文档表格中移除空白行"""
        for table in doc.tables:
            # 从后向前遍历表格行
            for row_idx in range(len(table.rows) - 1, -1, -1):
                row = table.rows[row_idx]
                is_blank_row = True
                
                # 检查行是否为空
                for cell in row.cells:
                    # 检查单元格中的段落
                    has_content = False
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            has_content = True
                            break
                    
                    # 如果单元格有表格嵌套，也需要检查
                    if not has_content and cell.tables:
                        has_content = True
                    
                    if has_content:
                        is_blank_row = False
                        break
                
                # 如果是空行，则删除
                if is_blank_row:
                    tbl = table._tbl
                    tr = row._tr
                    tbl.remove(tr)
                else:
                    # 如果不是空行，检查并删除单元格中的空白段落
                    for cell in row.cells:
                        # 从后向前遍历单元格段落
                        for para_idx in range(len(cell.paragraphs) - 1, -1, -1):
                            if self.is_blank_paragraph(cell.paragraphs[para_idx]):
                                p = cell.paragraphs[para_idx]._element
                                p.getparent().remove(p)
    
    def remove_section_blank_paragraphs(self, doc):
        """删除节中的空白段落"""
        for section in doc.sections:
            # 处理页眉
            for paragraph in section.header.paragraphs:
                if self.is_blank_paragraph(paragraph):
                    p = paragraph._element
                    p.getparent().remove(p)
            
            # 处理页脚
            for paragraph in section.footer.paragraphs:
                if self.is_blank_paragraph(paragraph):
                    p = paragraph._element
                    p.getparent().remove(p)
    
    def process_docx_file(self, input_path, output_path):
        """处理单个Word文件"""
        try:
            doc = Document(input_path)
            
            # 根据选项执行不同的清理操作
            if self.remove_paragraph_blanks.get():
                self.remove_paragraph_blank_lines(doc)
            
            if self.remove_table_blanks.get():
                self.remove_table_blank_lines(doc)
            
            if self.remove_section_blanks.get():
                self.remove_section_blank_paragraphs(doc)
            
            doc.save(output_path)
            return True, "成功"
        except Exception as e:
            return False, str(e)
    
    def start_processing(self):
        """开始处理所有Word文件"""
        input_dir = self.input_folder.get()
        output_dir = self.output_folder.get()
        
        if not input_dir or not output_dir:
            messagebox.showerror("错误", "请先选择输入和输出文件夹")
            return
        
        if not os.path.exists(input_dir):
            messagebox.showerror("错误", "输入文件夹不存在")
            return
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 获取所有Word文件
        word_files = [f for f in os.listdir(input_dir) 
                     if f.endswith(('.docx', '.doc')) and not f.startswith('~$')]
        
        if not word_files:
            messagebox.showinfo("信息", "输入文件夹中没有找到Word文档")
            return
        
        # 禁用按钮，开始处理
        self.status_label.config(text="处理中...")
        self.progress['maximum'] = len(word_files)
        self.progress['value'] = 0
        self.result_text.delete(1.0, tk.END)
        
        # 在新线程中处理，避免界面冻结
        thread = threading.Thread(target=self.process_files, args=(input_dir, output_dir, word_files))
        thread.daemon = True
        thread.start()
    
    def process_files(self, input_dir, output_dir, word_files):
        """处理文件线程函数"""
        success_count = 0
        fail_count = 0
        
        for i, filename in enumerate(word_files):
            input_path = os.path.join(input_dir, filename)
            output_path = os.path.join(output_dir, filename)
            
            # 更新进度
            self.root.after(0, self.update_progress, i, f"正在处理: {filename}")
            
            # 处理文件
            success, message = self.process_docx_file(input_path, output_path)
            
            if success:
                success_count += 1
                result_msg = f"✓ {filename}: 成功\n"
            else:
                fail_count += 1
                result_msg = f"✗ {filename}: 失败 - {message}\n"
            
            # 更新结果文本
            self.root.after(0, self.update_result, result_msg)
        
        # 完成处理
        final_msg = f"\n处理完成! 成功: {success_count}, 失败: {fail_count}"
        self.root.after(0, self.processing_complete, final_msg)
    
    def update_progress(self, value, status):
        """更新进度条和状态"""
        self.progress['value'] = value + 1
        self.status_label.config(text=status)
    
    def update_result(self, message):
        """更新结果文本框"""
        self.result_text.insert(tk.END, message)
        self.result_text.see(tk.END)
    
    def processing_complete(self, message):
        """处理完成"""
        self.status_label.config(text="处理完成")
        self.result_text.insert(tk.END, message)
        self.result_text.see(tk.END)
        messagebox.showinfo("完成", message)

if __name__ == "__main__":
    root = tk.Tk()
    app = WordBlankLineRemover(root)
    root.mainloop()